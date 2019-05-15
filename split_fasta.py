#!/usr/bin/python3
# the goal of this script is to take a fasta file from a normal gene, 
#and cut these into smaller parts (smaller than 100 aa or 300 nucleotides).
#Create a word file to put it in and than we can analyze the new sequences with the 
#coding algorithm. Check if these algorithm work perfectly

# import section

import os 
import glob
import re
from textwrap import wrap
from docx import Document

document = Document()
document.add_heading('Mathias Vanhockerhout:fasta splitting', 0)


# take file out of filesystem and open it

# look for each *.html file in current directory
#print ("this will look for .fa files in your home directory")
#location = "/home/mathias/"
#os.chdir(location)
#print("\nLook for .html files in {}:".format(location))
#for filename in glob.iglob('*.fa'):
#    print(filename)

print ("Welcome in this script that will criple fasta files")
loco = input("Geef de directory in waar je file zit")
location = ("{}".format(loco))
os.chdir(location)
print("\nLook for your file in {}:".format(location))
name_file = input("Hoe heet de file, vergeet niet om .fa ernaast te plaatsen")

for filename in glob.iglob(name_file):
    print(filename)


### wtf
full_loco = ("{}/{}".format(loco,name_file))
print ("\nHierna zou de full_loco moeten komen {}".format(full_loco))


fileObj = open(full_loco, "r")
c = 1
for line in fileObj.readlines():
    resultA = re.search("^A", line)
    resultT = re.search("^T", line)
    resultC = re.search("^C", line)
    resultG = re.search("^G", line)
    if resultA or resultT or resultC or resultG:
        chicken_wrap = wrap(line,100)
        document.add_paragraph('>fasta sequence from {}, sequence nr. {}'.format(name_file,c))
        document.add_paragraph(chicken_wrap)
    # Increase line counter
    c = c + 1
fileObj.close()

doc_name = input("Hoe wil je je document noemen, vergeet de .docx of .txt niet bij te vermelden")

answer = input("Just to confirm, you want to name your document: {}".format(doc_name))
if answer =='yes':
    document.save(doc_name)
    print ("Oke, your document is in your home directory")
else:
    print ("Ctr + C and start all over")




